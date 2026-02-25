"""
SimpleClaw v2.0 - File Generator Tool
=======================================
Generates files in multiple formats: PDF, DOCX, XLSX, CSV, PNG.
Files are saved to a temp directory and sent via Telegram.
"""

from __future__ import annotations

import csv
import io
import json
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

import structlog

logger = structlog.get_logger()

OUTPUT_DIR = Path("/tmp/simpleclaw_files")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def _output_path(filename: str) -> Path:
    """Generate timestamped output path."""
    stem = Path(filename).stem
    suffix = Path(filename).suffix
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return OUTPUT_DIR / f"{stem}_{timestamp}{suffix}"


# ─── CSV ────────────────────────────────────────────────────

def generate_csv(
    data: list[dict],
    filename: str = "output.csv",
    delimiter: str = ",",
) -> Path:
    """
    Generate a CSV file from a list of dicts.

    Args:
        data: List of dicts (each dict = one row)
        filename: Output filename
        delimiter: CSV delimiter

    Returns:
        Path to generated file
    """
    if not data:
        raise ValueError("Data list is empty")

    filepath = _output_path(filename)
    columns = list(data[0].keys())

    with open(filepath, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=columns, delimiter=delimiter)
        writer.writeheader()
        writer.writerows(data)

    logger.info("file.csv_generated", path=str(filepath), rows=len(data))
    return filepath


# ─── XLSX ───────────────────────────────────────────────────

def generate_xlsx(
    data: list[dict],
    filename: str = "output.xlsx",
    sheet_name: str = "Dados",
    header_style: bool = True,
) -> Path:
    """
    Generate an Excel file with formatted headers.

    Args:
        data: List of dicts
        filename: Output filename
        sheet_name: Worksheet name
        header_style: Apply bold/color to headers

    Returns:
        Path to generated file
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    if not data:
        raise ValueError("Data list is empty")

    filepath = _output_path(filename)
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name

    columns = list(data[0].keys())

    # Header row
    for col_idx, col_name in enumerate(columns, 1):
        cell = ws.cell(row=1, column=col_idx, value=col_name)
        if header_style:
            cell.font = Font(bold=True, color="FFFFFF", size=11)
            cell.fill = PatternFill(start_color="2F5496", end_color="2F5496", fill_type="solid")
            cell.alignment = Alignment(horizontal="center")

    # Data rows
    for row_idx, row_data in enumerate(data, 2):
        for col_idx, col_name in enumerate(columns, 1):
            ws.cell(row=row_idx, column=col_idx, value=row_data.get(col_name, ""))

    # Auto-adjust column widths
    for col_idx, col_name in enumerate(columns, 1):
        max_length = len(str(col_name))
        for row in data[:50]:  # Sample first 50 rows
            val = str(row.get(col_name, ""))
            max_length = max(max_length, len(val))
        ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = min(max_length + 2, 40)

    wb.save(filepath)
    logger.info("file.xlsx_generated", path=str(filepath), rows=len(data))
    return filepath


# ─── PDF ────────────────────────────────────────────────────

def generate_pdf(
    content: str,
    filename: str = "output.pdf",
    title: Optional[str] = None,
) -> Path:
    """
    Generate a PDF file from text content.

    Args:
        content: Text content for the PDF
        filename: Output filename
        title: Optional document title

    Returns:
        Path to generated file
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    filepath = _output_path(filename)
    doc = SimpleDocTemplate(
        str(filepath),
        pagesize=A4,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()
    story = []

    # Title
    if title:
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Title"],
            fontSize=18,
            spaceAfter=20,
        )
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 12))

    # Content - split by paragraphs
    body_style = ParagraphStyle(
        "CustomBody",
        parent=styles["Normal"],
        fontSize=11,
        leading=15,
        spaceAfter=8,
    )

    for paragraph in content.split("\n\n"):
        paragraph = paragraph.strip()
        if paragraph:
            # Escape XML special chars for ReportLab
            paragraph = (
                paragraph
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )
            story.append(Paragraph(paragraph, body_style))

    # Footer with timestamp
    footer_style = ParagraphStyle(
        "Footer",
        parent=styles["Normal"],
        fontSize=8,
        textColor="#888888",
    )
    story.append(Spacer(1, 30))
    story.append(Paragraph(
        f"Gerado por SimpleClaw em {datetime.now().strftime('%d/%m/%Y %H:%M')}",
        footer_style,
    ))

    doc.build(story)
    logger.info("file.pdf_generated", path=str(filepath))
    return filepath


# ─── DOCX ───────────────────────────────────────────────────

def generate_docx(
    content: str,
    filename: str = "output.docx",
    title: Optional[str] = None,
) -> Path:
    """
    Generate a Word document from text content.

    Args:
        content: Text content
        filename: Output filename
        title: Optional document title

    Returns:
        Path to generated file
    """
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    filepath = _output_path(filename)
    doc = Document()

    # Title
    if title:
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Content
    for paragraph in content.split("\n\n"):
        paragraph = paragraph.strip()
        if not paragraph:
            continue

        # Detect headers (lines starting with #)
        if paragraph.startswith("### "):
            doc.add_heading(paragraph[4:], level=3)
        elif paragraph.startswith("## "):
            doc.add_heading(paragraph[3:], level=2)
        elif paragraph.startswith("# "):
            doc.add_heading(paragraph[2:], level=1)
        else:
            p = doc.add_paragraph(paragraph)
            p.style.font.size = Pt(11)

    # Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph(
        f"Gerado por SimpleClaw em {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    )
    footer.style.font.size = Pt(8)

    doc.save(filepath)
    logger.info("file.docx_generated", path=str(filepath))
    return filepath


# ─── CHART / IMAGE ──────────────────────────────────────────

def generate_chart(
    data: dict,
    chart_type: str = "bar",
    filename: str = "chart.png",
    title: Optional[str] = None,
    xlabel: Optional[str] = None,
    ylabel: Optional[str] = None,
) -> Path:
    """
    Generate a chart image using matplotlib.

    Args:
        data: Dict with 'labels' and 'values' (or 'datasets' for multi-series)
        chart_type: bar, line, pie, scatter
        filename: Output filename
        title: Chart title
        xlabel: X-axis label
        ylabel: Y-axis label

    Returns:
        Path to generated file
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    filepath = _output_path(filename)
    fig, ax = plt.subplots(figsize=(10, 6))

    labels = data.get("labels", [])
    values = data.get("values", [])

    if chart_type == "bar":
        ax.bar(labels, values, color="#2F5496", edgecolor="white")
    elif chart_type == "line":
        ax.plot(labels, values, marker="o", color="#2F5496", linewidth=2)
    elif chart_type == "pie":
        ax.pie(values, labels=labels, autopct="%1.1f%%", startangle=90)
    elif chart_type == "scatter":
        x_vals = data.get("x", values)
        y_vals = data.get("y", values)
        ax.scatter(x_vals, y_vals, color="#2F5496", alpha=0.7)

    if title:
        ax.set_title(title, fontsize=14, fontweight="bold")
    if xlabel:
        ax.set_xlabel(xlabel)
    if ylabel:
        ax.set_ylabel(ylabel)

    plt.tight_layout()
    plt.savefig(filepath, dpi=150, bbox_inches="tight")
    plt.close()

    logger.info("file.chart_generated", path=str(filepath), type=chart_type)
    return filepath


# ─── CODE FILES ─────────────────────────────────────────────

def generate_code_file(
    content: str,
    filename: str = "script.py",
) -> Path:
    """Save code content to a file."""
    filepath = _output_path(filename)
    filepath.write_text(content, encoding="utf-8")
    logger.info("file.code_generated", path=str(filepath))
    return filepath
